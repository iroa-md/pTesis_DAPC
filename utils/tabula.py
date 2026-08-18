"""tabula.py — tablas paper-ready en formato Word.

Construye documentos a partir de tablas ya calculadas; no computa estadística. Emite
separador decimal de coma y agrupa las filas por bloque temporal mediante una banda de
sección, en el formato de las tablas del manuscrito.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field

import pandas as pd

from . import config, daten

IZQ = ("Variable", "Nivel", "Etiqueta", "Causa", "Predictor", "Especificación", "Algoritmo",
       "Estrato")

# Columnas cuyo valor se destaca en negrita cuando la fila alcanza significación.
COL_SIGNIFICACION = ("p", "p (vs desenlace)", "Test (p)")

# Identidad gráfica (config): reglas horizontales en azul hielo, texto en azul
# profundo, banda de sección en cian claro y acento coral solo para la
# significación estadística. Sin líneas verticales, al modo de las tablas
# académicas de tres líneas.
_HEX = lambda c: c.lstrip("#").upper()


@dataclass
class Spec:
    """Definición de una tabla: nombre de archivo, título y nota al pie."""
    stem: str
    titulo: str
    nota: str = ""
    grupo: str | None = None          # columna que agrupa las filas en bandas
    cols: list[str] = field(default_factory=list)
    anchos: dict | None = None        # solo LaTeX: fracción por columna, normalizada
    apaisada: bool = False            # solo LaTeX: gira la página, para tablas anchas
    etiqueta: str | None = None       # solo LaTeX: por omisión, «tab:» más el stem
    pagina_completa: bool = False     # solo LaTeX: reduce los márgenes en sus páginas


def coma(v: object) -> object:
    """Punto decimal a coma, preservando texto no numérico."""
    if isinstance(v, float):
        return f"{v:.2f}".replace(".", ",")
    if isinstance(v, str):
        return re.sub(r"(?<=\d)\.(?=\d)", ",", v)
    return v


def normalizar(df: pd.DataFrame, proteger: tuple = ()) -> pd.DataFrame:
    """Aplica el separador decimal a todas las columnas salvo las protegidas."""
    d = df.copy()
    for c in d.columns:
        if c not in proteger:
            d[c] = d[c].map(coma)
    return d


def _significativa(row: pd.Series) -> bool:
    """Verdadero si la fila alcanza significación estadística (p < 0,05)."""
    p = row.get("p")
    if p is None or (isinstance(p, float) and pd.isna(p)):
        return False
    t = str(p).strip().replace(",", ".")
    if t.startswith("<"):
        return True
    try:
        return float(t) < 0.05
    except ValueError:
        return False


_ESCAPES = {"\\": r"\textbackslash{}", "&": r"\&", "%": r"\%", "$": r"\$",
            "#": r"\#", "_": r"\_", "{": r"\{", "}": r"\}",
            "~": r"\textasciitilde{}", "^": r"\textasciicircum{}"}


def _tex(v: object) -> str:
    """Texto de celda listo para LaTeX.

    Las celdas de estas tablas traen «%» en los encabezados de intervalo y «&» en las
    etiquetas de nivel, que son justamente los dos caracteres que rompen una tabla.
    """
    s = "" if v is None or (isinstance(v, float) and pd.isna(v)) else str(v)
    return "".join(_ESCAPES.get(c, c) for c in s)


# Comandos LaTeX admitidos dentro del texto que se escapa. Es una lista blanca deliberada:
# todo lo demás se sigue escapando, de modo que una celda o una nota no puedan romper la
# tabla. `cref` resuelve secciones y figuras; `cite` y sus variantes, la bibliografía, que
# es lo que permite que una tabla de literatura numere por biblatex en vez de transcribir
# apellidos que quedarían desincronizados del documento.
_CRUDO = re.compile(r"\\(?:[cC]ref|cite|autocite|parencite|textcite)\{[^}]*\}")


def _tex_crudo(txt) -> str:
    """Texto escapado, salvo los comandos de la lista blanca, que pasan crudos.

    Sin esto una celda o una nota que cite imprime `\\textbackslash{}cite{...}` literal, que
    es lo que ocurrió con la nota de `tabla_M0_espinal`.

    Acepta cualquier valor y no solo texto, como ya hacía `_tex`: una celda numérica -la
    columna `n` de `tabla_exposicion_desenlace_centro`- llega como entero y hacía caer la
    emisión del fragmento LaTeX con `expected string or bytes-like object`, después de que el
    `.docx` de la misma tabla ya se había escrito.
    """
    txt = "" if txt is None or (isinstance(txt, float) and pd.isna(txt)) else str(txt)
    partes, ultimo = [], 0
    for m in _CRUDO.finditer(txt):
        partes += [_tex(txt[ultimo:m.start()]), m.group(0)]
        ultimo = m.end()
    partes.append(_tex(txt[ultimo:]))
    return "".join(partes)


# Compatibilidad con el nombre anterior, que solo cubría la nota al pie.
_tex_nota = _tex_crudo


def _nota_plana(txt: str) -> str:
    """Texto sin los comandos de la lista blanca, para Word, que no puede resolverlos.

    El `.docx` es un artefacto suelto, sin numeración de secciones ni bibliografía, de modo
    que dejar el comando visible sería peor que resolverlo a lo que pueda mostrarse. De una
    referencia cruzada no queda nada -y se retira el paréntesis que deja vacío-, y de una
    cita quedan sus claves, que es lo único trazable fuera del documento compilado.
    """
    t = re.sub(r"\\(?:cite|autocite|parencite|textcite)\{([^}]*)\}",
               lambda m: ", ".join(k.strip() for k in m.group(1).split(",")), txt)
    t = re.sub(r"\\[cC]ref\{[^}]*\}", "", t)
    t = re.sub(r"\s*\(\s*\)", "", t)
    return re.sub(r"\s{2,}", " ", t).strip()


def construir_tex(df: pd.DataFrame, sp: Spec, dest=None) -> None:
    """Escribe la tabla como fragmento LaTeX, para que el manuscrito la incluya.

    Emite un `longtable` de tres líneas al modo booktabs, con el encabezado repetido en
    cada página. La etiqueta permite referirla con `\\cref` sin escribir su número, de
    modo que reordenar las tablas no obligue a tocar el texto.

    Es la contraparte de `construir_docx` sobre el mismo DataFrame: una sola fuente y dos
    salidas, en lugar de transcribir la tabla al manuscrito.
    """
    cols = sp.cols or [c for c in df.columns if c != sp.grupo]
    dest = dest or (config.TABLAS / f"{sp.stem}.tex")
    dest.parent.mkdir(parents=True, exist_ok=True)
    etiqueta = sp.etiqueta or f"tab:{sp.stem.replace('tabla_', '')}"

    # Anchos proporcionales al contenido, y **todas** las columnas con ancho declarado.
    # Dejar alguna sin `p{}` la hace tomar su ancho natural, que no entra en el reparto y
    # desborda la página: es la forma habitual en que una tabla generada se vuelve
    # ilegible. El peso se recorta por arriba para que una celda muy larga no se coma el
    # espacio de las demás, y por abajo para que ninguna columna quede impracticable.
    largo = {c: max([len(str(c))] + [len(str(v)) for v in df[c]]) for c in cols}
    # `Spec.anchos` sustituye el reparto proporcional cuando el largo de la fuente no
    # predice el del papel. Es el caso de una columna de citas: en el marco mide cientos de
    # caracteres de claves y en la página imprime «[5, 10, 14]», de modo que el reparto
    # automático le da un ancho que no necesita y se lo quita a la columna de texto.
    peso = ({c: float(sp.anchos.get(c, 1)) for c in cols} if sp.anchos
            else {c: min(max(largo[c], 6), 60) for c in cols})
    # El espacio entre columnas no forma parte del ancho declarado y hay que descontarlo.
    disponible = max(0.40, 0.97 - 0.013 * len(cols))
    total = sum(peso.values()) or 1
    ancho = {c: disponible * peso[c] / total for c in cols}
    # Las columnas de texto se alinean a la izquierda; las de cifras, centradas.
    #
    # Las de texto van además en bandera y no justificadas. Un `p{}` justifica por omisión,
    # y en una columna estrecha eso reparte el blanco entre las pocas palabras que caben: un
    # «OR 1,71 [0,92-3,19]» de dos líneas sale con el rótulo pegado al borde izquierdo y la
    # cifra al derecho, como si fueran dos datos distintos. En bandera la celda se lee como
    # un solo valor partido en dos líneas, que es lo que es.
    texto = {c for c in cols if largo[c] > 18 or c in IZQ}
    align = "".join(
        (f">{{\\raggedright\\arraybackslash}}p{{{ancho[c]:.3f}\\linewidth}}" if c in texto
         else f">{{\\centering\\arraybackslash}}p{{{ancho[c]:.3f}\\linewidth}}")
        for c in cols)

    # La orientación y los márgenes NO se emiten aquí. Girar la página o cambiar la
    # geometría obliga a un salto, y si el fragmento lo emite, el título del anexo queda
    # solo en la página anterior. El envoltorio se declara en el markdown del anexo, donde
    # se ve qué comparte página con qué. `Spec.apaisada` y `Spec.pagina_completa`
    # sobreviven como documentación de lo que esa tabla necesita.
    L = []
    L += [r"\begingroup\footnotesize\setlength{\tabcolsep}{4pt}",
          r"\renewcommand{\arraystretch}{1.15}",
          r"\begin{longtable}{" + align + "}",
          r"\caption{" + _tex(sp.titulo) + r"}\label{" + etiqueta + r"}\\",
          r"\toprule",
          " & ".join(r"\textbf{" + _tex(c) + "}" for c in cols) + r" \\",
          r"\midrule", r"\endfirsthead",
          r"\toprule",
          " & ".join(r"\textbf{" + _tex(c) + "}" for c in cols) + r" \\",
          r"\midrule", r"\endhead",
          r"\bottomrule", r"\endlastfoot"]

    grupo_actual = None
    for _, row in df.iterrows():
        if sp.grupo and row[sp.grupo] != grupo_actual:
            grupo_actual = row[sp.grupo]
            L.append(r"\multicolumn{" + str(len(cols)) + r"}{l}{\textit{"
                     + _tex(grupo_actual) + r"}} \\")
        # Las celdas pasan por la lista blanca, no por el escape ciego: una tabla de
        # literatura lleva `\cite` en su columna de referencias.
        L.append(" & ".join(_tex_crudo(row[c]) for c in cols) + r" \\")

    L += [r"\end{longtable}", r"\endgroup"]
    if sp.nota:
        L.append(r"\noindent\footnotesize " + _tex_nota(sp.nota) + r"\normalsize")

    dest.write_text("\n".join(L) + "\n", encoding="utf-8")
    envoltorio = ("apaisada" if (sp.apaisada or len(cols) > 6)
                  else ("página completa" if sp.pagina_completa else ""))
    print(f"✔ {dest.name}  ({len(df)} filas)"
          + (f"  → requiere {envoltorio} en el anexo" if envoltorio else ""))


def construir_docx(df: pd.DataFrame, sp: Spec, dest=None) -> None:
    """Escribe la tabla en un documento de Word apaisado con su título y nota.

    Emite además el fragmento LaTeX equivalente, porque el manuscrito se compila desde
    LaTeX y una tabla que solo existe en Word no llega al documento.
    """
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    cols = sp.cols or [c for c in df.columns if c != sp.grupo]
    dest = dest or (config.TABLAS / f"{sp.stem}.docx")
    dest.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, m, Pt(48))

    def _rgb(c):
        h = c.lstrip("#")
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    tinta, tenue = _rgb(config.COLOR_TEXTO), _rgb(config.GRIS_TENUE)
    regla = _HEX(config.COLOR_MEDIO)
    h = doc.add_paragraph().add_run(sp.titulo)
    h.bold, h.font.size, h.font.name, h.font.color.rgb = True, Pt(11), config.TIPOGRAFIA[0], tinta

    tbl = doc.add_table(rows=1, cols=len(cols))
    tbl.alignment, tbl.autofit = WD_TABLE_ALIGNMENT.CENTER, True

    def borde(celda, lados):
        tcPr = celda._tc.get_or_add_tcPr()
        b = OxmlElement("w:tcBorders")
        for lado, (sz, color) in lados.items():
            el = OxmlElement(f"w:{lado}")
            el.set(qn("w:val"), "single"); el.set(qn("w:sz"), str(sz)); el.set(qn("w:color"), color)
            b.append(el)
        tcPr.append(b)

    def sombra(celda, hexcolor):
        tcPr = celda._tc.get_or_add_tcPr()
        sh = OxmlElement("w:shd")
        sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), _HEX(hexcolor))
        tcPr.append(sh)

    def fmt(celda, txt, *, bold=False, size=8, color=tinta, centrar=False):
        celda.text = ""
        p = celda.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centrar else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(2)
        run = p.add_run("" if pd.isna(txt) else _nota_plana(str(txt)))
        run.bold, run.font.size, run.font.name, run.font.color.rgb = bold, Pt(size), config.TIPOGRAFIA[0], color

    hdr = tbl.rows[0]
    for i, c in enumerate(cols):
        fmt(hdr.cells[i], c, bold=True, centrar=c not in IZQ)
        borde(hdr.cells[i], {"top": (12, _HEX(config.COLOR_TEXTO)), "bottom": (6, regla)})
    trPr = hdr._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader"); th.set(qn("w:val"), "true"); trPr.append(th)

    grupo_actual = None
    for _, row in df.iterrows():
        if sp.grupo and row[sp.grupo] != grupo_actual:
            grupo_actual = row[sp.grupo]
            fila = tbl.add_row()
            fila.cells[0].merge(fila.cells[len(cols) - 1])
            fmt(fila.cells[0], grupo_actual, bold=True)
            sombra(fila.cells[0], config.COLOR_FONDO)
            borde(fila.cells[0], {"top": (6, regla)})
        fila = tbl.add_row()
        destacar = _significativa(row)
        for i, c in enumerate(cols):
            fmt(fila.cells[i], row[c], centrar=c not in IZQ,
                bold=destacar and c in COL_SIGNIFICACION)

    for celda in tbl.rows[-1].cells:
        borde(celda, {"bottom": (12, _HEX(config.COLOR_TEXTO))})

    if sp.nota:
        n = doc.add_paragraph().add_run(_nota_plana(sp.nota))
        n.font.size, n.font.name, n.font.color.rgb = Pt(7), config.TIPOGRAFIA[0], tenue

    doc.save(dest)
    print(f"✔ {dest.name}  ({len(df)} filas)")
    construir_tex(df, sp)


# ─────────────────────────────────────────────────────────────────────────────
# Constructores de las tablas del manuscrito
# ─────────────────────────────────────────────────────────────────────────────

ETIQUETA_BLOQUE = {"preop": "Preoperatorio", "intraop": "Intraoperatorio",
                   "postop": "Postoperatorio temprano"}


def _p_fmt(v) -> str:
    """Valor p con coma decimal; por debajo de 0,001 se reporta como umbral."""
    if v is None or (isinstance(v, float) and pd.isna(v)):
        return ""
    return "< 0,001" if v < 0.001 else f"{v:.3f}".replace(".", ",")


def tabla_descriptiva(d: pd.DataFrame) -> pd.DataFrame:
    """Tabla maestra de características y asociación cruda con el desenlace.

    El marco de `vista` viene en dos clases de fila: una por variable, con sus estadísticos,
    y bajo las de más de dos niveles una por categoría, con solo su reparto. Aquí ambas se
    funden en una sola columna, con las categorías precedidas de un punto medio: una columna
    «Nivel» propia queda vacía en toda fila de variable, y estrecha además la de etiquetas,
    que es la que más ancho necesita.

    La medida cruda se rotula con su estadístico. Sin eso, la única de las cuatro que no se
    identifica sola es el odds ratio -`r`, `V` y `d` ya vienen prefijados-, y la columna
    mezcla cuatro escalas distintas sin decir cuál es cuál en cada fila.
    """
    t = d.copy()
    t["Bloque"] = t["bloque"].map(ETIQUETA_BLOQUE).fillna(t["bloque"])
    t["p (vs desenlace)"] = t["p"].map(_p_fmt)
    t["n análisis"] = t["n análisis"].map(lambda v: "" if v in ("", None) else str(int(v)))
    t["Variable"] = [f"· {niv}" if str(niv).strip() else var
                     for var, niv in zip(t["Variable"], t["Nivel"])]
    # El OR llega sin rótulo, como «1,78 [1,25-2,52]»: se reconoce porque empieza por dígito
    # y trae intervalo. Es un rótulo, no una transformación de la cifra.
    t["Medida cruda"] = [f"OR {m}" if re.match(r"^\d.*\[", str(m)) else m
                         for m in t["Medida cruda"]]
    cols = ["Bloque", "Variable", "Global", "Faltantes", "n análisis",
            "Prueba", "p (vs desenlace)", "Medida cruda", "p"]
    return normalizar(t[cols], proteger=("Bloque", "Variable", "Faltantes",
                                         "n análisis", "Prueba", "p (vs desenlace)", "p"))


def tabla_por_centro(d: pd.DataFrame) -> pd.DataFrame:
    """Tabla de características por centro y heterogeneidad entre hospitales."""
    t = d.copy()
    t["Bloque"] = t["bloque"].map(ETIQUETA_BLOQUE).fillna(t["bloque"])
    centros = [c for c in t.columns if c.endswith(")") and "n=" in c]
    cols = ["Bloque", "Variable", "Faltantes", *centros, "Test (p)", "SMD máx. (centro)", "p"]
    return normalizar(t[cols], proteger=("Bloque", "Variable", "Test (p)",
                                         "SMD máx. (centro)", "p"))


def tabla_sesgo_seleccion(d: pd.DataFrame) -> pd.DataFrame:
    """Anexo de sesgo de selección: incluidas frente a cada grupo de excluidas."""
    t = d.copy()
    t["Grupo"] = t["estrato"] + " (n=" + t["n_excluidas"].astype(str) + ")"
    t["p"] = t["p"].map(_p_fmt)
    t["SMD"] = t["SMD"].map(lambda v: "" if pd.isna(v) else f"{v:.3f}".replace(".", ","))
    inc = next(c for c in t.columns if c.startswith("Incluidas"))
    # Sin la columna «Frágil»: su contenido se explica en la nota al pie.
    cols = ["Grupo", "Variable", inc, "Excluidas", "SMD", "Prueba", "p"]
    return normalizar(t[cols], proteger=tuple(cols))


def tabla_excluidas(excluidas: pd.DataFrame, sp: Spec) -> pd.DataFrame:
    """Anexo de variables excluidas del análisis y su causa."""
    d = excluidas.copy()
    d["Bloque"] = d["block"].map(ETIQUETA_BLOQUE).fillna(d["block"])
    causa = d["exclusion_causa"].astype(str)
    detalle = d["exclusion_nota"].fillna("").astype(str)
    d["Causa"] = [c if not x.strip() else x for c, x in zip(causa, detalle)]
    d = d.rename(columns={"label": "Variable"})
    return normalizar(d[["Bloque", "Variable", "Causa"]], proteger=("Bloque", "Variable"))


def tabla_frecuencia_severidad(d: pd.DataFrame) -> pd.DataFrame:
    """Frecuencia de dolor por momento de medición, con la proporción con dolor como resumen.

    `d` trae las cuatro categorías excluyentes de `vista.frecuencia_severidad_tiempo` (la
    misma que arma la figura). «Con dolor» no es una quinta categoría del agregado, es
    leve + moderado + intenso, y se deriva aquí porque es de presentación, no de cálculo:
    los recuentos que suma ya vienen calculados. No se llama «incidencia»: ese término se
    reserva para el desenlace principal (`dmg_1s_bin`), no para un corte transversal por
    momento.
    """
    etiqueta_rango = {
        "Leve": f"Leve (END 1-{daten.UMBRAL_MG - 1})",
        "Moderado": f"Moderado (END {daten.UMBRAL_MG}-{daten.UMBRAL_GRAVE - 1})",
        "Intenso": f"Intenso (END {daten.UMBRAL_GRAVE}-10)",
    }
    resumen = "Con dolor (END > 0)"
    filas = []
    for momento, g in d.groupby("momento", sort=False, observed=True):
        g = g.set_index("categoria")
        n_total = int(g["n_total"].iloc[0])
        n_dolor = int(g.loc[list(etiqueta_rango), "n"].sum())
        filas.append({"Categoría": resumen, "momento": momento,
                      "celda": f"{n_dolor}/{n_total} ({100 * n_dolor / n_total:.1f}%)"})
        for cat, etq in etiqueta_rango.items():
            n = int(g.loc[cat, "n"])
            filas.append({"Categoría": etq, "momento": momento,
                          "celda": f"{n}/{n_total} ({100 * n / n_total:.1f}%)"})
    t = pd.DataFrame(filas)
    orden = [resumen, *etiqueta_rango.values()]
    t["Categoría"] = pd.Categorical(t["Categoría"], orden, ordered=True)
    ancho = t.pivot(index="Categoría", columns="momento", values="celda")
    ancho.columns.name = None
    ancho = ancho.reset_index()
    cols = ["Categoría", "24 horas", "48 horas", "7 días"]
    return normalizar(ancho[cols], proteger=("Categoría",))


def tabla_frecuencia_centro_tiempo(d: pd.DataFrame) -> pd.DataFrame:
    """Frecuencia de dolor por centro y momento de medición, con IC 95 % de Wilson.

    No se llama «incidencia»: ese término se reserva para el desenlace principal
    (`dmg_1s_bin`), no para un corte transversal por momento.
    """
    t = d.copy()
    t["celda"] = [f"{100 * inc:.1f} % [{100 * lo:.1f} – {100 * hi:.1f}]"
                 for inc, lo, hi in zip(t["incidencia"], t["ic_low"], t["ic_high"])]
    ancho = t.pivot(index="centro", columns="momento", values="celda")
    ancho.columns.name = None
    orden = [*sorted(c for c in ancho.index if c != "GLOBAL"), "GLOBAL"]
    ancho = ancho.reindex(orden).reset_index().rename(columns={"centro": "Centro"})
    # "GLOBAL" es el valor interno de statistik.frecuencia_por_centro; en la tabla
    # publicada, como en el texto del manuscrito, va en minúscula salvo la inicial.
    ancho["Centro"] = ancho["Centro"].replace({"GLOBAL": "Global"})
    cols = ["Centro", "24 horas", "48 horas", "7 días"]
    return normalizar(ancho[cols], proteger=("Centro",))


# ── Arco explicativo ─────────────────────────────────────────────────────────

def _ic(low, high) -> str:
    """Intervalo de confianza en el formato del manuscrito."""
    if pd.isna(low) or pd.isna(high):
        return ""
    return f"{low:.2f} – {high:.2f}".replace(".", ",")


def tabla_modelo_explicativo(t: pd.DataFrame, cat: pd.DataFrame) -> pd.DataFrame:
    """Términos del modelo explicativo con su odds ratio ajustado."""
    meta = cat.set_index("var_rename")
    d = t[t["termino"].ne("Intercept")].copy()
    d["bloque"] = d["var_rename"].map(lambda v: meta.loc[v, "block"] if v in meta.index else "")
    d["Bloque"] = d["bloque"].map(ETIQUETA_BLOQUE).fillna(d["bloque"])
    d["Predictor"] = d["label"]
    d["OR ajustado"] = d["OR"].map(lambda v: f"{v:.2f}".replace(".", ","))
    d["IC 95 %"] = [_ic(a, b) for a, b in zip(d["ic_low"], d["ic_high"])]
    d["p"] = d["p"].map(_p_fmt)
    d = d.sort_values("bloque", key=lambda s: s.map(config.ORDEN_BLOQUE).fillna(9))
    cols = ["Bloque", "Predictor", "OR ajustado", "IC 95 %", "p"]
    return normalizar(d[cols], proteger=tuple(cols))


def tabla_escalera(esc: pd.DataFrame) -> pd.DataFrame:
    """Escalera de modelos anidados y aporte de cada bloque de covariables."""
    d = esc.copy()
    d["Peldaño"] = d["peldano"]
    d["Covariables incorporadas"] = d["descripcion"]
    d["Parámetros"] = d["n_parametros"].astype(int)
    d["EPV"] = d["EPV"].map(lambda v: f"{v:.1f}".replace(".", ","))
    for c in ("AIC", "BIC"):
        d[c] = d[c].map(lambda v: f"{v:.1f}".replace(".", ","))
    d["R² McFadden"] = d["R2_McFadden"].map(lambda v: f"{v:.3f}".replace(".", ","))
    d["Razón de verosimilitud (gl)"] = [
        "" if pd.isna(v) else f"{v:.2f} ({int(g)})".replace(".", ",")
        for v, g in zip(d.get("lrt", pd.Series(dtype=float)), d.get("gl", pd.Series(dtype=float)))]
    d["p"] = d.get("p_lrt", pd.Series(dtype=float)).map(_p_fmt)
    # Sin AIC, BIC ni R²: la comparación entre peldaños la resuelve la razón de
    # verosimilitud, y los tres índices de ajuste no se interpretan en el texto.
    cols = ["Peldaño", "Covariables incorporadas", "Parámetros", "EPV",
            "Razón de verosimilitud (gl)", "p"]
    return normalizar(d[cols], proteger=tuple(cols))


def tabla_vif(v: pd.DataFrame) -> pd.DataFrame:
    """Colinealidad por variable, en su forma generalizada."""
    d = v.copy()
    d["Variable"] = d["label"]
    d["Grados de libertad"] = d["gl"].astype(int)
    d["GVIF"] = d["GVIF"].map(lambda x: f"{x:.3f}".replace(".", ","))
    d["GVIF^(1/2gl)"] = d["GVIF_escalado"].map(lambda x: f"{x:.3f}".replace(".", ","))
    d["VIF equivalente"] = d["VIF_equivalente"].map(lambda x: f"{x:.2f}".replace(".", ","))
    # Se reporta solo el VIF equivalente, que es la cifra interpretable frente al umbral.
    # El GVIF y su raíz quedan fuera: son el paso intermedio del cálculo.
    cols = ["Variable", "Grados de libertad", "VIF equivalente"]
    return normalizar(d[cols], proteger=tuple(cols))


def tabla_estabilidad(e: pd.DataFrame) -> pd.DataFrame:
    """Frecuencia de selección de cada variable bajo remuestreo penalizado."""
    d = e.copy()
    d["Variable"] = d["label"]
    d["Frecuencia de selección"] = d["frecuencia"].map(lambda v: f"{v:.2f}".replace(".", ","))
    d["Signo dominante"] = d["signo"]
    d["Consistencia del signo"] = d["consistencia_signo"].map(
        lambda v: "" if pd.isna(v) else f"{v:.2f}".replace(".", ","))
    d["Estabilidad"] = d["estabilidad"]
    cols = ["Variable", "Frecuencia de selección", "Signo dominante",
            "Consistencia del signo", "Estabilidad"]
    return normalizar(d[cols], proteger=tuple(cols))


def tabla_por_desenlace(ajustes: dict, cat: pd.DataFrame, *, etiquetas: dict | None = None,
                        medidas: dict | None = None) -> pd.DataFrame:
    """Estimaciones del mismo modelo bajo varias definiciones del desenlace, en paralelo.

    Una fila por término y una columna por definición, de modo que la fila completa muestre
    si la asociación se sostiene al cambiar la operacionalización. Sustituye a una tabla por
    desenlace: el contraste entre definiciones es precisamente lo que se quiere leer, y en
    tablas separadas obliga a saltar de página. La significación se marca con un asterisco,
    porque una columna por definición no deja espacio para el valor p.
    """
    etiquetas, medidas = etiquetas or {}, medidas or {}
    meta = cat.set_index("var_rename")
    base = next(iter(ajustes.values()))
    orden = base[base["termino"].ne("Intercept") & ~base["umbral"]]

    filas = []
    for _, r in orden.iterrows():
        v = r["var_rename"]
        bloque = meta.loc[v, "block"] if v in meta.index else ""
        fila = {"bloque": bloque, "Bloque": ETIQUETA_BLOQUE.get(bloque, bloque),
                "Predictor": r["label"]}
        for des, t in ajustes.items():
            m = t[t["termino"].eq(r["termino"])]
            if m.empty:
                fila[etiquetas.get(des, des)] = ""
                continue
            e = m.iloc[0]
            cel = (f"{e['efecto']:.2f} [{e['ic_low']:.2f}–{e['ic_high']:.2f}]"
                   .replace(".", ","))
            fila[etiquetas.get(des, des)] = cel + ("*" if e["p"] < 0.05 else "")
        filas.append(fila)

    d = pd.DataFrame(filas).sort_values(
        "bloque", key=lambda s: s.map(config.ORDEN_BLOQUE).fillna(9), kind="stable")
    cols = ["Bloque", "Predictor"] + [etiquetas.get(k, k) for k in ajustes]
    return d[cols]


def tabla_apoyo_seleccion(bivariado: pd.DataFrame, estabilidad: pd.DataFrame,
                          cat: pd.DataFrame) -> pd.DataFrame:
    """Comportamiento de cada predictor candidato en los dos criterios de apoyo.

    Documenta la asociación cruda y la frecuencia de selección bajo remuestreo penalizado
    junto a la decisión adoptada, de modo que se vea cuándo el criterio clínico primó sobre
    ellos. Ambos criterios operan subordinados: pueden respaldar la incorporación de un
    predictor no previsto, no excluir un factor de riesgo reportado en la literatura.
    """
    meta = cat.set_index("var_rename")
    ROL = {"literatura": "Incluida · factor de riesgo reportado",
           "centro": "Incluida · ajuste por centro",
           "ajustador": "Incluida · ajustador clínico",
           "bivariado_estabilidad": "Incluida · asociación bivariada y estabilidad"}

    d = estabilidad.merge(bivariado[["var_rename", "p"]], on="var_rename", how="left")
    filas = []
    for _, r in d.iterrows():
        v = r["var_rename"]
        m = meta.loc[v] if v in meta.index else {}
        rol = m.get("rol_explicativo")
        causa = m.get("exclusion_causa")
        if isinstance(rol, str) and rol in ROL:
            decision = ROL[rol]
        elif isinstance(causa, str) and causa.strip():
            decision = f"Excluida · {causa.strip()}"
        elif bool(m.get("include_model", False)):
            # Candidata del arco predictivo: no entra al conjunto parsimonioso porque este
            # se define por los factores de riesgo reportados, no por el desempeño en los datos.
            decision = "No incluida en el conjunto parsimonioso · candidata del arco predictivo"
        else:
            decision = "No incluida"
        bloque = m.get("block", "")
        filas.append({"bloque": bloque, "Bloque": ETIQUETA_BLOQUE.get(bloque, bloque),
                      "Predictor": r["label"], "p bivariado": _p_fmt(r.get("p")),
                      "Frec. LASSO": f"{r['frecuencia']:.2f}".replace(".", ","),
                      "Estabilidad": str(r["estabilidad"]).capitalize(),
                      "Decisión": decision})
    d = pd.DataFrame(filas).sort_values(
        ["bloque", "Frec. LASSO"], key=lambda s: (s.map(config.ORDEN_BLOQUE).fillna(9)
                                                  if s.name == "bloque" else s),
        ascending=[True, False])
    # Sin «Decisión»: repite lo que la sección ya declara y es la columna más ancha.
    cols = ["Bloque", "Predictor", "p bivariado", "Frec. LASSO", "Estabilidad"]
    return d[cols]


def tabla_epv(filas: list[dict]) -> pd.DataFrame:
    """Eventos por parámetro de cada modelo, frente al criterio orientador."""
    d = pd.DataFrame(filas)
    d["EPV"] = d["EPV"].map(lambda v: f"{v:.1f}".replace(".", ","))
    d["Cumple"] = d["cumple"].map({True: "Sí", False: "No"})
    d = d.rename(columns={"modelo": "Modelo", "n_parametros": "Parámetros",
                          "eventos": "Eventos", "n": "n"})
    return d[["Modelo", "n", "Parámetros", "Eventos", "EPV", "Cumple"]]


def tabla_desempeno(d: pd.DataFrame, *, escenario: str = "sin_centro") -> pd.DataFrame:
    """Discriminación y calibración de cada algoritmo en las dos especificaciones.

    Las filas se agrupan por especificación, de modo que la comparación que responde a la
    hipótesis se lea por columnas y no reconstruyendo el orden.
    """
    d = d[d["escenario"].eq(escenario)].copy()
    d["Especificación"] = d["especificacion"].map(config.NOMBRE_ESPECIFICACION)
    d["Algoritmo"] = d["algoritmo"]
    d["AUC-ROC (IC95 %)"] = [f"{a:.3f} ({lo:.3f} – {hi:.3f})".replace(".", ",")
                             for a, lo, hi in zip(d["auc"], d["auc_ic_low"], d["auc_ic_high"])]
    d["Brier"] = d["brier"].map(lambda v: f"{v:.3f}".replace(".", ","))
    d["Pendiente de calibración"] = d["pendiente"].map(lambda v: f"{v:.2f}".replace(".", ","))
    if "citl" in d.columns:
        d["Calibración media"] = d["citl"].map(lambda v: f"{v:.3f}".replace(".", ","))
    cols = ["Especificación", "Algoritmo", "AUC-ROC (IC95 %)", "Brier",
            "Pendiente de calibración"] + (["Calibración media"] if "citl" in d.columns else [])
    d = d.sort_values(["especificacion", "auc"], ascending=[True, False])
    return normalizar(d[cols], proteger=tuple(cols))


def tabla_pareada(p: pd.DataFrame, *, etiquetas: dict | None = None,
                  escenarios: dict | None = None) -> pd.DataFrame:
    """Diferencia pareada entre dos especificaciones, por algoritmo.

    Cada fila compara un algoritmo consigo mismo sobre las mismas pacientes, de modo que la
    diferencia no arrastre variabilidad entre muestras. Se reportan ambas métricas, porque
    una mejora del ordenamiento no implica una mejora de las probabilidades.

    `escenarios` traduce el escenario que codifica la clave de `referencia` a la etiqueta con
    que debe encabezar su bloque. Sin él, un marco que contrasta la misma pareja de
    especificaciones bajo dos escenarios produce dos series de filas idénticas en todo salvo
    en las cifras, y nada en la tabla dice cuál es cuál: es lo que ocurría al contrastar la
    hipótesis con y sin la variable centro. La columna resultante se declara como `grupo` en
    el `Spec`, que la convierte en banda de sección y la retira de las columnas visibles.
    """
    etiquetas, escenarios = etiquetas or {}, escenarios or {}
    d = p.copy()
    d["Contraste"] = [etiquetas.get(r, r) for r in d["referencia"]]
    d["Algoritmo"] = d["alternativa"].str.rsplit("·", n=1).str[-1].str.strip()
    if escenarios:
        # La clave viaja como «cohorte · escenario · especificación · algoritmo».
        clave = d["referencia"].str.split("·").str[1].str.strip()
        d["Escenario"] = [escenarios.get(e, e) for e in clave]
    d["ΔAUC-ROC (IC95 %)"] = [f"{v:.3f} ({lo:.3f} a {hi:.3f})".replace(".", ",")
                              for v, lo, hi in zip(d["delta_auc"], d["delta_auc_ic_low"],
                                                   d["delta_auc_ic_high"])]
    d["p"] = d["p_unilateral"].map(_p_fmt)
    d["ΔBrier (IC95 %)"] = [f"{v:.3f} ({lo:.3f} a {hi:.3f})".replace(".", ",")
                            for v, lo, hi in zip(d["delta_brier"], d["delta_brier_ic_low"],
                                                 d["delta_brier_ic_high"])]
    d["p "] = d["p_brier_unilateral"].map(_p_fmt)
    cols = (["Escenario"] if escenarios else []) + \
           ["Contraste", "Algoritmo", "ΔAUC-ROC (IC95 %)", "p", "ΔBrier (IC95 %)", "p "]
    return normalizar(d[cols], proteger=tuple(cols))


ETIQUETA_ROL = {"literatura": "Factor de riesgo reportado en la literatura",
                "bivariado_estabilidad": "Asociación bivariada y estabilidad en la selección",
                "ajustador": "Incorporada como ajustador",
                "centro": "Unidad de agrupamiento del estudio"}


def tabla_predictores_asociaciones(cat: pd.DataFrame) -> pd.DataFrame:
    """Predictores del modelo de asociaciones, con su bloque y el criterio que los incorpora.

    Declara por qué entró cada uno, que es lo que separa una selección justificada de una
    lista. El criterio proviene del catálogo y no de la memoria de quien escribe.
    """
    d = cat[cat["include_explicativo"].fillna(False)].copy()
    d["Bloque"] = d["block"].map(ETIQUETA_BLOQUE).fillna(d["block"])
    d["Predictor"] = d["label"]
    d["Tipo"] = d["conceptual_type"].map(
        {"num": "Numérica", "bin": "Binaria", "ord": "Ordinal", "nom": "Nominal"}).fillna("")
    d["Nivel de referencia"] = d["positive_label"].where(
        d["conceptual_type"].ne("bin"), "No").fillna("")
    d["Criterio de inclusión"] = d["rol_explicativo"].map(ETIQUETA_ROL).fillna("")
    orden = {v: i for i, v in enumerate(ETIQUETA_BLOQUE.values())}
    d = d.sort_values(["Bloque", "Predictor"], key=lambda s: s.map(orden).fillna(99)
                      if s.name == "Bloque" else s)
    # Sin «Nivel de referencia»: ocupa una columna entera para decir «No» en las binarias.
    cols = ["Bloque", "Predictor", "Tipo", "Criterio de inclusión"]
    return normalizar(d[cols], proteger=tuple(cols))


def tabla_especificaciones(cat: pd.DataFrame, predictores, cohorte: str = "principal") -> pd.DataFrame:
    """Composición de las especificaciones preoperatoria y perioperatoria.

    Una fila por predictor candidato y una marca por especificación, de modo que la
    diferencia entre ambas, que es lo que contrasta la hipótesis, se lea de un vistazo en
    lugar de comparar dos listas.

    `predictores` es `modell.predictores`, que se recibe como argumento para no acoplar este
    módulo al de modelamiento.
    """
    meta = cat.set_index("var_rename")
    conjuntos = {e: set(predictores(cat, e, "sin_centro", cohorte)) for e in ("preop", "periop")}
    filas = []
    for v in sorted(conjuntos["periop"] | conjuntos["preop"]):
        b = str(meta.loc[v, "block"]) if v in meta.index else ""
        filas.append({"bloque": b,
                      "Bloque": ETIQUETA_BLOQUE.get(b, b),
                      "Predictor": str(meta.loc[v, "label"]) if v in meta.index else v,
                      "Preoperatoria": "Sí" if v in conjuntos["preop"] else "—",
                      "Perioperatoria": "Sí" if v in conjuntos["periop"] else "—"})
    orden = {k: i for i, k in enumerate(ETIQUETA_BLOQUE)}
    d = pd.DataFrame(filas).sort_values(["bloque", "Predictor"],
                                        key=lambda s: s.map(orden) if s.name == "bloque" else s)
    cols = ["Bloque", "Predictor", "Preoperatoria", "Perioperatoria"]
    return normalizar(d[cols], proteger=tuple(cols))


def tabla_hiperparametros(hp: pd.DataFrame, grillas: dict) -> pd.DataFrame:
    """Rejilla explorada por algoritmo y valores elegidos en los pliegues externos.

    La selección ocurre dentro de cada pliegue, de modo que un algoritmo puede terminar con
    valores distintos según el pliegue. Se reporta el recuento de cada valor elegido, que
    informa si la elección fue estable o no. Se separa por especificación, porque el mismo
    algoritmo puede preferir otra penalización cuando dispone de más predictores.
    """
    filas = []
    for esp in [e for e in config.ESPECIFICACIONES if e in set(hp.get("especificacion", []))] or [None]:
        g_esp = hp if esp is None else hp[hp["especificacion"].eq(esp)]
        nombre = config.NOMBRE_ESPECIFICACION.get(esp, "") if esp else ""
        for alg in [a for a in config.ALGORITMOS if a in set(g_esp["algoritmo"])]:
            g = g_esp[g_esp["algoritmo"].eq(alg)]
            rejilla = grillas.get(alg) or {}
            for par, gp in g.groupby("hiperparametro", sort=False):
                if par == "—":
                    filas.append({"Especificación": nombre, "Algoritmo": alg,
                                  "Hiperparámetro": "—", "Rejilla explorada": "—",
                                  "Elegidos por pliegue": "—"})
                    continue
                vals = gp["valor"].astype(str).value_counts()
                explorada = rejilla.get(par, [])
                filas.append({
                    "Especificación": nombre, "Algoritmo": alg, "Hiperparámetro": par,
                    "Rejilla explorada": ", ".join(str(x) for x in explorada) if len(explorada) else "—",
                    "Elegidos por pliegue": ", ".join(f"{k} ({v})" for k, v in vals.items())})
    cols = ["Especificación", "Algoritmo", "Hiperparámetro", "Rejilla explorada",
            "Elegidos por pliegue"]
    return normalizar(pd.DataFrame(filas)[cols],
                      proteger=("Especificación", "Algoritmo", "Hiperparámetro"))


def tabla_transportabilidad(iecv: pd.DataFrame) -> pd.DataFrame:
    """AUC-ROC de cada algoritmo en el centro que quedó fuera de su entrenamiento.

    Una columna por centro retenido y su promedio, que es la cifra que resume la caída. Se
    ordena por promedio descendente, de modo que el mejor y el peor queden en los extremos.
    """
    d = iecv.pivot_table(index="algoritmo", columns="centro", values="auc")
    d["Promedio"] = d.mean(axis=1)
    d = d.sort_values("Promedio", ascending=False).reset_index()
    d = d.rename(columns={"algoritmo": "Algoritmo"})
    for c in d.columns:
        if c != "Algoritmo":
            d[c] = d[c].map(lambda v: "" if pd.isna(v) else f"{v:.3f}".replace(".", ","))
    return normalizar(d, proteger=tuple(d.columns))


def tabla_logistica_comparada(ajustes: dict, cat: pd.DataFrame, *,
                              solo: list | None = None) -> pd.DataFrame:
    """Odds ratio de una misma especificación estimada en dos cohortes, lado a lado.

    Sirve para leer la dirección de las asociaciones de los predictores que la importancia
    por permutación sitúa arriba, sin constituir un segundo análisis de factores. Las
    columnas se alinean por término, de modo que un predictor ausente en una de las dos
    quede visible como tal y no desplace las filas.

    `solo` restringe las filas a un conjunto de variables, que es lo habitual cuando la
    tabla acompaña a un consenso de importancia y no pretende reportar el modelo completo.
    """
    orden, filas = [], {}
    for nombre, t in ajustes.items():
        d = t[t["termino"].ne("Intercept")]
        if solo is not None:
            d = d[d["var_rename"].isin(solo)]
        for _, r in d.iterrows():
            k = r["termino"]
            if k not in filas:
                filas[k] = {"Predictor": str(r["label"])}
                orden.append(k)
            filas[k][f"{nombre}, OR (IC95 %)"] = (
                f"{r['OR']:.2f}".replace(".", ",") + f" ({_ic(r['ic_low'], r['ic_high'])})"
                if pd.notna(r["OR"]) else "")
            filas[k][f"{nombre}, p"] = _p_fmt(r["p"])
    cols = ["Predictor"] + [c for n in ajustes for c in (f"{n}, OR (IC95 %)", f"{n}, p")]
    d = pd.DataFrame([filas[k] for k in orden]).reindex(columns=cols).fillna("no incluido")
    # Sin proteger: las etiquetas de nivel arrastran el punto decimal de la fuente («7.5 mg»)
    # y `coma` solo sustituye entre dígitos, de modo que no toca lo ya formateado.
    return normalizar(d)


def tabla_logistica_consenso(ajuste: pd.DataFrame, puestos: dict) -> pd.DataFrame:
    """Odds ratio de una cohorte, restringidos y ordenados por el consenso de importancia.

    Es la contraparte de una sola cohorte de `tabla_logistica_comparada`. La versión cruzada
    tomaba la unión de los predictores destacados en ambas, y como las especificaciones
    difieren -la espinal retira la analgesia postoperatoria y recodifica el manejo
    intraoperatorio- la mitad de las filas quedaba a medio poblar, con «no incluido» en una
    de las dos columnas. Una tabla por cohorte, con el conjunto de esa cohorte, no tiene esa
    celda.

    `puestos` es el mapa `var_rename → puesto en el consenso` y hace las dos cosas a la vez:
    **selecciona** las filas y **las ordena**. Con eso la tabla se lee fila a fila contra la
    figura de consenso y contra los paneles SHAP de su misma cohorte, en lugar de por bloque
    temporal o por orden alfabético, que obligan a buscar cada predictor. El puesto se
    imprime en su propia columna para que ese orden sea verificable y no haya que confiar en
    él.

    Una nominal aporta un término por nivel, de modo que puede dar más de una fila: se
    ordenan por puesto y, dentro de él, en el orden en que el diseño las produjo. El nivel se
    anexa al rótulo, porque sin él dos filas del mismo predictor se imprimen idénticas.
    """
    d = ajuste[ajuste["termino"].ne("Intercept") & ajuste["var_rename"].isin(puestos)].copy()
    d["_puesto"] = d["var_rename"].map(puestos)
    d = d.sort_values("_puesto", kind="stable")

    t = pd.DataFrame({
        "Puesto": d["_puesto"].astype(int).astype(str).to_numpy(),
        "Predictor": [f"{lab} · {niv}" if pd.notna(niv) and str(niv) else str(lab)
                      for lab, niv in zip(d["label"], d["nivel"])],
        "OR (IC95 %)": [f"{o:.2f}".replace(".", ",") + f" ({_ic(lo, hi)})" if pd.notna(o) else ""
                        for o, lo, hi in zip(d["OR"], d["ic_low"], d["ic_high"])],
        "p": [_p_fmt(p) for p in d["p"]]})
    # Sin `normalizar`: las celdas ya llevan su coma y las etiquetas de nivel arrastran el
    # punto decimal de la fuente («7,5 mg»), que `coma` volvería a tocar.
    return t.reset_index(drop=True)


def tabla_exposicion_desenlace_centro(datos: pd.DataFrame, exposicion: str, *,
                                      desenlace: str | None = None, centro: str | None = None,
                                      rotulos: dict | None = None) -> pd.DataFrame:
    """Frecuencia de una exposición y del desenlace, lado a lado, por conglomerado.

    Hace legible en una sola lectura si ambas corren en la misma dirección entre centros o
    en direcciones opuestas, que es lo que distingue una asociación clínica de una que el
    ajuste por conglomerado produce.

    **El denominador es el total del centro, no los casos observados.** Con faltantes en la
    exposición las dos convenciones difieren, y esta es la que ya publica el manuscrito. La
    consecuencia -una fila con exposición desconocida se cuenta en el denominador pero no en
    el numerador- obliga a declarar los faltantes en la nota, y por eso se devuelven en
    `attrs["faltantes"]` en vez de dejarlos implícitos.

    Los recuentos se calculan aquí sobre el dato y no se transcriben, que es lo que permite
    que la tabla contradiga a quien la invoca si la cifra citada dejó de ser cierta.
    """
    centro = centro or config.CENTER_VAR
    desenlace = desenlace or config.OUTCOME
    rotulos = rotulos or {}
    celda = lambda k, n: f"{int(k)}/{int(n)} ({100 * k / n:.1f} %)".replace(".", ",")

    filas = []
    for c in sorted(datos[centro].dropna().unique()):
        d = datos[datos[centro].eq(c)]
        filas.append({"Centro": str(c), "n": len(d),
                      rotulos.get("exposicion", "Exposición"): celda(d[exposicion].sum(), len(d)),
                      rotulos.get("desenlace", "Desenlace"): celda(d[desenlace].sum(), len(d))})
    filas.append({"Centro": "Total", "n": len(datos),
                  rotulos.get("exposicion", "Exposición"):
                      celda(datos[exposicion].sum(), len(datos)),
                  rotulos.get("desenlace", "Desenlace"):
                      celda(datos[desenlace].sum(), len(datos))})

    t = pd.DataFrame(filas)
    t.attrs["faltantes"] = {"exposicion": int(datos[exposicion].isna().sum()),
                            "desenlace": int(datos[desenlace].isna().sum())}
    # Sin `normalizar`: las celdas ya llevan la coma decimal y su barra n/N no debe tocarse.
    return t


def tabla_desenlace_por_estrato_centro(d: pd.DataFrame) -> pd.DataFrame:
    """Desenlace dentro de cada estrato de exposición, por centro, en bandas por predictor.

    El agregado lo produce `vista.desenlace_por_estrato_centro`. Aquí solo se rotulan las
    columnas y se agrupa: no se generaliza `tabla_exposicion_desenlace_centro`, que tiene
    los ejes traspuestos -una fila por centro, no una por estrato- y mide frecuencias
    marginales del centro en vez del desenlace dentro del estrato. Un solo nombre para las
    dos obligaría a un interruptor que cambia la forma de la salida, que son dos funciones
    con un nombre.
    """
    t = d.rename(columns={"predictor": "Predictor", "estrato": "Estrato"})
    cols = ["Predictor", "Estrato"] + [c for c in t.columns
                                       if c not in ("Predictor", "Estrato")]
    # Las celdas de cifra ya traen su coma decimal y su barra n/N no debe tocarse. «Estrato»
    # es la excepción: sus rótulos son los niveles tal como los declara el catálogo, con
    # punto decimal -«< 7.5 mg»-, porque son la llave de unión con el dato. `coma` los pasa
    # a la ortografía del manuscrito sin que haya que mantener un segundo juego de rótulos.
    return normalizar(t[cols], proteger=tuple(c for c in cols if c != "Estrato"))


def tabla_conversiones(espinal: pd.DataFrame, cat: pd.DataFrame,
                       conversiones: list[dict]) -> pd.DataFrame:
    """Recodificaciones derivadas, con su origen, su regla y su distribución observada.

    La distribución se recuenta sobre el dato, no se transcribe: es la comprobación de que
    la regla declarada produce lo que la tabla afirma.

    El reparto se escribe sin abreviaturas ni signos: la tabla va al manuscrito y sus celdas
    se leen como texto, no como notación.
    """
    meta = cat.set_index("var_rename")
    filas = []
    for c in conversiones:
        v = c["var_rename"]
        s = espinal[v] if v in espinal.columns else pd.Series(dtype=object)
        if s.dtype.name in ("bool", "boolean"):
            partes = [f"Sí {int(s.sum())}", f"No {int((s == False).sum())}"]
        else:
            partes = [f"{k} {n}" for k, n in
                      s.value_counts(dropna=True).reindex(
                          s.cat.categories if hasattr(s, "cat") else None
                      ).dropna().astype(int).items()]
        faltan = int(s.isna().sum())
        if faltan:
            partes.append(f"sin dato {faltan}")
        fila = {"Recodificada": str(meta.loc[v, "label"]) if v in meta.index else v,
                "Variables de origen": c["origen"],
                "Regla de conversión": c["regla"],
                "Distribución observada": ", ".join(partes)}
        # No todas las recodificaciones alimentan a los dos arcos, y cuál entra a cuál es
        # justamente lo que la tabla debe dejar explícito.
        if "uso" in c:
            fila["Análisis en que se emplea"] = c["uso"]
        filas.append(fila)
    return pd.DataFrame(filas)
